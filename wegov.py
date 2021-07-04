from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
import sys, os

# ファイル情報
pyName = 'WeGov'
pyVersion = '0.2'
pyAuthor = 'FTS141'
pyComments = 'https://github.com/fts141/WeGov'


def xmlTreeExploration(soup, level=1):
    for content in soup.contents:
        if content.name is None or content.name == 'TOC': continue
        level = writeExportFile(content, level)
        if len(content) > 1:
            xmlTreeExploration(content, level)


def writeExportFile(soup, level=1):
    
    name = soup.name

    if name == 'LawNum':
        global lawNum
        lawNum = soup.text

    elif name == 'LawTitle':
        global lawTitle
        lawTitle = soup.text
        exportFile.add_heading('{}\n（{}）'.format(soup.text, lawNum), level=1)

    elif name == 'SupplProvisionLabel':
        amendLawNum = soup.parent.get('AmendLawNum')
        if amendLawNum is not None:
            writeText = '{}（{}）'.format(soup.text.replace('　',''), amendLawNum)
        else:
            writeText = '{}'.format(soup.text)
        exportFile.add_page_break()
        exportFile.add_heading('{}'.format(writeText), level=1)

    elif name == 'ChapterTitle':
        exportFile.add_heading(soup.text, level=2)

    elif name == 'SectionTitle':
        exportFile.add_heading(soup.text, level=3)

    elif name == 'ArticleTitle':
        articleCaption = soup.parent.ArticleCaption
        if articleCaption is not None:
            writeText = '{}{}'.format(soup.text, articleCaption.text)
        else:
            writeText = soup.text
        exportFile.add_heading(writeText, level=4)

    elif name == 'ParagraphSentence':
        paragraphNum = soup.parent.ParagraphNum.text
        if paragraphNum != '':
            exportFile.add_heading(paragraphNum, level=5).paragraph_format.left_indent = indent(1)
            exportFile.add_paragraph(soup.text.strip()).paragraph_format.left_indent = indent(1)
        else:
            exportFile.add_paragraph(soup.text.strip())

    elif name == 'ItemSentence':
        itemTitle = soup.parent.ItemTitle.text
        exportFile.add_paragraph( \
            '{}\n{}'.format(itemTitle, soup.text.strip()), style='List Bullet').paragraph_format.left_indent = indent(2)

    else: pass
    return level


def writeDocumentProperties():
    global exportFile

    properties = exportFile.core_properties
    properties.title = u'{}（{}）'.format(lawTitle, lawNum)
    properties.author = u'{}'.format(pyAuthor)
    properties.comments = u'{}'.format(pyComments)
    properties.language = u'ja-JP'      # 機能していない？


def writeHeaderAndFooter():
    global exportFile

    header = exportFile.sections[0].header.paragraphs[0]
    header.text = '\t{}\n\t（{}）'.format(lawTitle, lawNum)
    header.style = exportFile.styles['Header']

    footer = exportFile.sections[0].footer.paragraphs[0]
    footer.text = '\t{} v{} - {}'.format(pyName, pyVersion, pyAuthor)
    footer.style = exportFile.styles['Footer']


def indent(indentLevel):
    span =  Pt(11) * 2      # インデント幅（文字サイズ * 文字数）
    return span * indentLevel


if __name__ == '__main__':

    lawNum = ''
    lawTitle = ''

    try:
        soup = BeautifulSoup(open(sys.argv[1], 'r', encoding='UTF-8'), 'lxml-xml')
        exportDirectory = '{}/'.format(sys.argv[2])
    except:
        print('''
        <!> 実行できませんでした。引数を確認してください。
        実行例 -> python3 wegov.py importFile.xml ./exportDir
        ''')
        exit()

    exportFile = Document()
    xmlTreeExploration(soup.Law)
    writeDocumentProperties()
    writeHeaderAndFooter()

    exportFileName = '{}（{}）.docx'.format(lawTitle, lawNum)
    exportDirectory = '{}/{}/'.format(exportDirectory, lawTitle)
    if not os.path.exists(exportDirectory): os.makedirs(exportDirectory)
    
    try:
        exportFile.save('{}{}'.format(exportDirectory, exportFileName))
        print('''
        (i) ファイルを書き出しました。 -> {}
        '''.format(exportFileName))
    except:
        print('''
        <!> ファイルの書き込みに失敗しました。同名のファイルが開かれていませんか？
        ''')
